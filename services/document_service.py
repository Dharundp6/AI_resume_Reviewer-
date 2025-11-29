"""
Document Generation Service
Creates Word documents (.docx) for resume and cover letter
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime


class DocumentService:
    """Service for generating Word documents"""
    
    @staticmethod
    def create_resume_docx(content: str, company_name: str, output_dir: str = "outputs") -> str:
        """
        Create a Word document for the optimized resume
        
        Args:
            content: Resume text content
            company_name: Name of the company
            output_dir: Directory to save the file
            
        Returns:
            Path to the generated file
        """
        try:
            # Create a new Document
            doc = Document()
            
            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.75)
                section.bottom_margin = Inches(0.75)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)
            
            # Split content into paragraphs and add to document
            paragraphs = content.split('\n')
            
            for para_text in paragraphs:
                if para_text.strip():
                    # Detect if this is a header (all caps or starts with specific keywords)
                    is_header = (
                        para_text.isupper() or 
                        any(para_text.strip().startswith(word) for word in 
                            ['PROFESSIONAL', 'EDUCATION', 'EXPERIENCE', 'SKILLS', 
                             'PROJECTS', 'CERTIFICATIONS', 'SUMMARY'])
                    )
                    
                    paragraph = doc.add_paragraph(para_text)
                    
                    if is_header:
                        # Format as header
                        run = paragraph.runs[0]
                        run.bold = True
                        run.font.size = Pt(12)
                    else:
                        # Regular text
                        run = paragraph.runs[0] if paragraph.runs else None
                        if run:
                            run.font.size = Pt(11)
            
            # Generate filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{company_name.replace(' ', '_')}_Optimized_Resume_{timestamp}.docx"
            filepath = os.path.join(output_dir, filename)
            
            # Ensure output directory exists
            os.makedirs(output_dir, exist_ok=True)
            
            # Save document
            doc.save(filepath)
            
            return filepath
            
        except Exception as e:
            raise Exception(f"Failed to create resume document: {str(e)}")
    
    @staticmethod
    def create_cover_letter_docx(content: str, company_name: str, output_dir: str = "outputs") -> str:
        """
        Create a Word document for the cover letter
        
        Args:
            content: Cover letter text content
            company_name: Name of the company
            output_dir: Directory to save the file
            
        Returns:
            Path to the generated file
        """
        try:
            # Create a new Document
            doc = Document()
            
            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Split content into paragraphs
            paragraphs = content.split('\n')
            
            for para_text in paragraphs:
                if para_text.strip():
                    paragraph = doc.add_paragraph(para_text)
                    
                    # Set font
                    for run in paragraph.runs:
                        run.font.size = Pt(11)
                        run.font.name = 'Calibri'
            
            # Generate filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{company_name.replace(' ', '_')}_Cover_Letter_{timestamp}.docx"
            filepath = os.path.join(output_dir, filename)
            
            # Ensure output directory exists
            os.makedirs(output_dir, exist_ok=True)
            
            # Save document
            doc.save(filepath)
            
            return filepath
            
        except Exception as e:
            raise Exception(f"Failed to create cover letter document: {str(e)}")


# Create singleton instance
document_service = DocumentService()
